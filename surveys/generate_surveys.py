# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_question(doc, num, text, qtype=None):
    label = f"{num}、{text}"
    if qtype == "multi":
        label += "【多选题】"
    elif qtype == "text":
        label += "【填空题】"
    p = doc.add_paragraph()
    run = p.add_run(label)
    set_font(run, bold=True, size=12)
    return p


def add_option(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run(f"{letter}、{text}")
    set_font(run, size=11)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(f"■ {title}")
    set_font(run, bold=True, size=13, color=(70, 130, 180))


def add_blank(doc):
    doc.add_paragraph()


# ───────────────────────── 问卷A：客人版 ─────────────────────────
def build_customer_survey():
    doc = Document()

    t = doc.add_heading("", level=0)
    r = t.add_run("档期预约平台用户调研问卷（客人版）")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d = doc.add_paragraph("感谢您参与本次调研！本问卷完全匿名，结果仅用于产品设计参考，填写约需 3 分钟。")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_blank(doc)

    # 第一部分
    add_section(doc, "第一部分：基本情况")

    add_question(doc, 1, "你通常约的服务类型是？", "multi")
    for l, o in zip("ABCD", ["妆娘", "摄影", "毛娘", "其他"]):
        add_option(doc, l, o)

    add_question(doc, 2, "你平均多久约一次服务？")
    for l, o in zip("ABC", ["每月一次以上", "每 2-3 个月一次", "半年以下"]):
        add_option(doc, l, o)

    add_question(doc, 3, "你现在通过什么渠道联系商家？", "multi")
    for l, o in zip("ABCDE", ["微博私信", "小红书私信", "微信", "QQ", "其他"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第二部分
    add_section(doc, "第二部分：现有痛点")

    add_question(doc, 4, "约档期时你遇到过哪些麻烦？", "multi")
    for l, o in zip("ABCDEF", [
        "不知道商家还有没有空档期，要私信问",
        "商家回复慢，确认档期要等很久",
        "付定金流程麻烦或没有保障",
        "档期临时变动，商家没有及时通知我",
        "不知道去哪里找合适的商家",
        "其他",
    ]):
        add_option(doc, l, o)

    add_question(doc, 5, "付定金时你通常用什么方式？")
    for l, o in zip("ABCD", ["支付宝转账", "微信转账", "闲鱼下单", "不付定金"]):
        add_option(doc, l, o)

    add_question(doc, 6, "你有没有遇到过付了定金但出问题的情况（跑单、档期冲突等）？")
    for l, o in zip("ABC", ["有过", "没有", "没付过定金"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第三部分
    add_section(doc, "第三部分：功能需求")

    add_question(doc, 7, "以下功能你觉得哪些最重要？（请选出最重要的 3 个）", "multi")
    for l, o in zip("ABCDEFG", [
        "直接查看商家空闲档期，不用私信询问",
        "固定时间开抢档期，大家公平排队",
        "选好档期后直接跳转支付宝/闲鱼付定金",
        "订单状态实时通知（商家确认后提醒我）",
        "看商家作品集和历史案例",
        "和商家在平台上直接沟通",
        "查看自己的历史订单记录",
    ]):
        add_option(doc, l, o)

    add_question(doc, 8, "你能接受的定金比例范围？")
    for l, o in zip("ABCD", ["10-20%", "20-30%", "30-50%", "按商家规定来"]):
        add_option(doc, l, o)

    add_question(doc, 9, "如果需要注册新平台，你倒向用哪种方式登录？")
    for l, o in zip("ABC", ["手机号验证码", "微信扫码", "微博账号"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第四部分
    add_section(doc, "第四部分：使用意愿")

    add_question(doc, 10, "如果这个平台上线且完全免费，你愿意使用吗？")
    for l, o in zip("ABCD", [
        "非常愿意",
        "愿意，先看看",
        "看有没有我认识的商家入驻",
        "不太愿意",
    ]):
        add_option(doc, l, o)

    add_question(doc, 11, "你还希望平台有什么功能？（开放填写）", "text")

    doc.save("E:/dyh/MHP/surveys/问卷A_客人版.docx")
    print("OK: 问卷A_客人版.docx")


# ───────────────────────── 问卷B：商家版 ─────────────────────────
def build_merchant_survey():
    doc = Document()

    t = doc.add_heading("", level=0)
    r = t.add_run("档期预约平台用户调研问卷（商家版）")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d = doc.add_paragraph("感谢您参与本次调研！本问卷完全匿名，结果仅用于产品设计参考，填写约需 3 分钟。")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_blank(doc)

    # 第一部分
    add_section(doc, "第一部分：基本情况")

    add_question(doc, 1, "你提供的服务类型？", "multi")
    for l, o in zip("ABC", ["妆娘", "摄影", "毛娘"]):
        add_option(doc, l, o)

    add_question(doc, 2, "你目前每月接单数量大概是？")
    for l, o in zip("ABC", ["5 单以下", "5-15 单", "15 单以上"]):
        add_option(doc, l, o)

    add_question(doc, 3, "你目前怎么管理档期？")
    for l, o in zip("ABCD", [
        "全靠记忆或聊天记录",
        "手机备忘录/日历",
        "表格",
        "没有专门管理",
    ]):
        add_option(doc, l, o)

    add_question(doc, 4, "你目前收定金的方式？", "multi")
    for l, o in zip("ABCD", ["支付宝转账", "微信转账", "闲鱼下单", "不收定金"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第二部分
    add_section(doc, "第二部分：现有痛点")

    add_question(doc, 5, "管理档期时你最大的烦恼是？", "multi")
    for l, o in zip("ABCDEFG", [
        "同一时间被多人私信，难以按先后顺序处理",
        "客人付定金不规范，容易跑单",
        "反复确认时间，沟通成本高",
        "档期信息要手动在各平台同步",
        "临时取消订单时定金纠纷难处理",
        "展示作品麻烦，各平台规则不同",
        "其他",
    ]):
        add_option(doc, l, o)

    add_question(doc, 6, "你有没有因为档期管理混乱导致撞期或漏单的情况？")
    for l, o in zip("ABC", ["有过，比较困扪", "偶尔有", "基本没有"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第三部分
    add_section(doc, "第三部分：功能需求")

    add_question(doc, 7, "以下功能哪些对你最重要？（请选出最重要的 3 个）", "multi")
    for l, o in zip("ABCDEFG", [
        "可视化日历管理档期（直观看到哪天有约）",
        "客人按先后顺序自动排队，我逐个联系",
        "客人选好档期后直接跳转我的支付宝/闲鱼付定金",
        "新订单自动通知我",
        "发布作品帖子展示例图",
        "客人可以在我主页直接看到空闲档期",
        "订单记录和历史存档",
    ]):
        add_option(doc, l, o)

    add_question(doc, 8, "关于“抢档期”功能（设定时间点，到点客人同时开抢，按服务器时间排序），这个模式适合你吗？")
    for l, o in zip("ABCD", [
        "非常适合，想用这个模式",
        "适合，但我也想保留私信约单的方式",
        "不太适合，我更喜欢自己筛选客人",
        "不了解，需要看看实际效果",
    ]):
        add_option(doc, l, o)

    add_question(doc, 9, "客人确认订单后，平台提供按键直接跳转到你的支付宝收款页或闲鱼商品页，你能接受这个付款方式吗？")
    for l, o in zip("ABCD", [
        "完全可以，我现在就是这样收款的",
        "可以接受",
        "希望平台能直接帮我收款（自动确认）",
        "无所谓",
    ]):
        add_option(doc, l, o)

    add_question(doc, 10, "你希望平台支持哪些跳转方式？", "multi")
    for l, o in zip("ABC", ["支付宝收款链接", "闲鱼商品链接", "两种都要"]):
        add_option(doc, l, o)

    add_blank(doc)

    # 第四部分
    add_section(doc, "第四部分：使用意愿")

    add_question(doc, 11, "如果这个平台完全免费，你愿意入驻吗？")
    for l, o in zip("ABCD", [
        "非常愿意",
        "愿意，先试试",
        "要看有没有客人在用",
        "不太愿意",
    ]):
        add_option(doc, l, o)

    add_question(doc, 12, "你觉得什么情况会让你放弃使用这个平台？", "multi")
    for l, o in zip("ABCDE", [
        "客人不够多",
        "操作太麻烦",
        "功能不满足需求",
        "平台开始收费",
        "其他",
    ]):
        add_option(doc, l, o)

    add_question(doc, 13, "你还希望平台有什么功能？（开放填写）", "text")

    doc.save("E:/dyh/MHP/surveys/问卷B_商家版.docx")
    print("OK: 问卷B_商家版.docx")


if __name__ == "__main__":
    build_customer_survey()
    build_merchant_survey()
    print("两份问卷均已生成。")
